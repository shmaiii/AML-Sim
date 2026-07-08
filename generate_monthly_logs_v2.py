"""Generate May and June 2026 monthly logs following the NUS template."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from copy import deepcopy

OUT_DIR = Path.home() / "Desktop"
TEMPLATE_PATH = Path.home() / "Desktop" / "Monthly Log template.docx"


def _copy_run_format(src, dst, text: str, bold=None) -> None:
    """Copy formatting from src run to dst run, optionally overriding bold."""
    dst.text = text
    if src.font.size:
        dst.font.size = src.font.size
    if src.font.name:
        dst.font.name = src.font.name
    if bold is not None:
        dst.bold = bold
    else:
        dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline


def _add_filled_paragraph(doc, template_para, replacements, *, bold_label=True):
    """
    Clone a template paragraph, replacing {placeholders} with actual values.
    If bold_label is True, the label part stays bold and the value part is not.
    """
    p = doc.add_paragraph()
    p.alignment = template_para.alignment
    p.paragraph_format.space_after = template_para.paragraph_format.space_after

    for run in template_para.runs:
        text = run.text
        for key, value in replacements.items():
            text = text.replace(key, value)
        r = p.add_run(text)
        _copy_run_format(run, r, text, bold=run.bold)
    return p


def _add_section_heading(doc, title: str) -> None:
    """Add a section heading like 'Achievement during the month:'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(title)
    r.font.size = Pt(11)
    r.font.name = "Calibri"


def _add_bullet(doc, label: str, body: str) -> None:
    """Add a bullet point with bold label and normal body text."""
    p = doc.add_paragraph(style="List Bullet")
    r_label = p.add_run(label + "  ")
    r_label.bold = True
    r_label.font.size = Pt(11)
    r_label.font.name = "Calibri"
    r_body = p.add_run(body)
    r_body.font.size = Pt(11)
    r_body.font.name = "Calibri"


def generate_monthly_log(month: str, year: str, period: str, content: dict) -> Document:
    """Generate a monthly log doc following the NUS template structure."""

    # Load template to clone its header paragraphs
    template = Document(str(TEMPLATE_PATH))

    doc = Document()

    # Page setup — match template
    section = doc.sections[0]
    section.page_width = template.sections[0].page_width
    section.page_height = template.sections[0].page_height
    section.left_margin = template.sections[0].left_margin
    section.right_margin = template.sections[0].right_margin
    section.top_margin = template.sections[0].top_margin
    section.bottom_margin = template.sections[0].bottom_margin

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # --- Header lines (cloned from template) ---
    _add_filled_paragraph(doc, template.paragraphs[0], {})  # NUS School of Computing
    _add_filled_paragraph(doc, template.paragraphs[1], {})  # FT5007 Capstone
    _add_filled_paragraph(doc, template.paragraphs[2], {
        "<   >": f"< {month} >",
    })  # Month < May > Report by Individual Team Member

    doc.add_paragraph()  # blank line

    # --- Project Title ---
    _add_filled_paragraph(doc, template.paragraphs[4], {
        "______________________": "AML-Sim: A Multi-Agent Market Simulation Platform for Studying Adaptive Decision-Making Under Uncertainty",
    })

    # --- Reporting Period (built from scratch — template runs are fragmented) ---
    rp = doc.add_paragraph()
    r1 = rp.add_run("Reporting Period: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = rp.add_run(period)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    r2.underline = True

    doc.add_paragraph()  # blank line

    # --- Instruction ---
    _add_filled_paragraph(doc, template.paragraphs[7], {})

    doc.add_paragraph()  # blank line

    # --- Student Name ---
    _add_filled_paragraph(doc, template.paragraphs[9], {
        "____________________": "PAN YUFAN",
    })

    # =========================================================================
    # Section 1: Achievements
    # =========================================================================
    _add_section_heading(doc, "Achievement during the month:")
    for label, body in content.get("achievements", []):
        _add_bullet(doc, label, body)

    doc.add_paragraph()  # spacer

    # =========================================================================
    # Section 2: To-dos
    # =========================================================================
    _add_section_heading(doc, "To-dos in next month:")
    for label, body in content.get("todos", []):
        _add_bullet(doc, label, body)

    doc.add_paragraph()  # spacer

    # =========================================================================
    # Section 3: Challenges
    # =========================================================================
    _add_section_heading(doc, "Challenges Faced:")
    for label, body in content.get("challenges", []):
        _add_bullet(doc, label, body)

    doc.add_paragraph()  # spacer

    # =========================================================================
    # Section 4: Plans to overcome
    # =========================================================================
    _add_section_heading(doc, "Plans to overcome challenges:")
    for label, body in content.get("plans", []):
        _add_bullet(doc, label, body)

    return doc


# =============================================================================
# MAY Content
# =============================================================================
may_content = {
    "achievements": [
        (
            "Literature review: Multi-agent LLM frameworks for financial tasks.",
            "Studied three key works that directly informed the AML-Sim design. "
            "(1) FinVision (Fatemi & Hu, arXiv 2411.08899v1): a multi-modal "
            "multi-agent framework for stock market prediction that deploys "
            "specialised LLM agents (Summarise, Technical Analyst, Prediction, "
            "and Reflection modules) with visual reasoning on candlestick charts "
            "and a reflection module that analyses historical trading signals and "
            "their outcomes. This paper demonstrated the value of a dedicated "
            "reflection/improvement loop — a concept we adopted in AML-Sim's "
            "slow-loop strategist. (2) Tan Shannon's B.Sc. Dissertation (NUS "
            "SOC, 2025/2026, supervised by Assoc Prof Anand Bhojan): Multi-Agent "
            "LLM Refinement Loops for RL-Driven Formulaic Alpha Discovery. "
            "This work introduces a Syntax-Fix LLM Agent that autonomously "
            "resolves parsing errors in generated alpha expressions, plus a "
            "dual-agent semantic loop where a Critique Agent analyses standalone "
            "Information Coefficient and pool correlation to direct an Improvement "
            "Agent in executing targeted structural modifications. The multi-agent "
            "refinement architecture — and its limitations around in-sample "
            "overfitting — directly shaped our validator design and the separation "
            "of strategy proposal from strategy application. (3) Chen & Kawashima "
            "(University of Hyogo): a two-agent LLM framework where DeepSeek "
            "generates a broad pool of formulaic alphas and Gemini evaluates and "
            "selects the top ten based on predictive strength and signal diversity. "
            "This generation-evaluation split validated our architecture decision "
            "to separate the slow strategist (proposer) from the validator (evaluator)."
        ),
        (
            "Reviewed the AML project brief and defined system scope.",
            "Analysed the AML Project Details document which decomposes the "
            "platform into four integrated systems: (1) Market Ecosystem "
            "Simulator — exchange, order book, liquidity, volatility; (2) AI "
            "Multi-Agent Environment — heterogeneous agents with different "
            "information, goals, latency, and risk appetite; (3) Controlled "
            "Uncertainty Lab — inject fake news, volatility shocks, information "
            "delays, liquidity crashes, adversarial spoofing, and policy "
            "interventions; and (4) Human Decision Experiment Platform — humans "
            "participate alongside AI agents. The core research question: 'What "
            "makes a decision-maker robust under uncertainty?' — measuring "
            "resilience, adaptability, stability, calibration, and recoverability "
            "rather than raw PnL."
        ),
        (
            "System architecture design: two-layer separation and dual-loop agents.",
            "Decided on a two-layer architecture where AML-Sim owns agent "
            "behaviour, strategy, memory, observation packaging, and experiment "
            "orchestration, while StockSim (integrated as a git submodule) owns "
            "execution, messaging, portfolio accounting, and order state. Designed "
            "the fast-loop / slow-loop dual-loop agent architecture inspired by "
            "the reflection mechanisms in FinVision and the refinement loops in "
            "Tan Shannon's work: a high-frequency execution loop (role-specific "
            "order placement) driven by a validated strategy state, and a lower-"
            "frequency strategic loop that proposes strategy updates (LLM-driven "
            "in future) — with all proposals passing through a bounds validator "
            "before application. Also defined the three initial agent roles: "
            "Market Maker (continuous bid/ask quoting), Retail Trader (noisy "
            "probabilistic orders), and Institutional Trader (sliced target "
            "execution)."
        ),
        (
            "Technology stack selection and environment setup.",
            "Confirmed the technology stack based on StockSim compatibility and "
            "project requirements: Python 3.13 for all components, RabbitMQ for "
            "inter-process pub/sub messaging, multiprocessing.Process for "
            "component isolation (each exchange, trader, and the clock runs in "
            "its own process), asyncio for agent-internal concurrency, YAML for "
            "declarative scenario definition, JSON for action ledgers and reports, "
            "and git submodules for managing the StockSim dependency. Set up the "
            "development environment on Windows 11 and integrated StockSim as a "
            "git submodule under simulators/StockSim."
        ),
        (
            "Proof-of-concept implementation: scenario runner and synthetic market.",
            "Wrote the initial aml_runner.py that reads an AML scenario YAML file, "
            "extracts and validates the required stocksim_config section (exchange "
            "mode, instruments, exchanges, agents, simulation parameters), creates "
            "timestamped run directories under .aml_runs/, and archives the "
            "original scenario, generated StockSim config, and run metadata for "
            "reproducibility. Created the first synthetic order-book scenario "
            "(aml_orderbook_replay.yaml) — a 30-minute AAPL simulation with one "
            "market maker (fair price $100, spread $0.20, quote size 100), five "
            "retail traders (45% trade probability, max size 10), and one "
            "institutional trader (target 500 shares, child order size 50) — "
            "using synthetic data so no external API keys are required."
        ),
        (
            "Prepared infrastructure presentation.",
            "Created and delivered the AML Infrastructure Presentation covering "
            "the project motivation, four-system architecture, technology stack, "
            "development roadmap, and the Systems + Infrastructure Lead role "
            "scope: exchange engine, architecture, concurrency, networking, "
            "simulation orchestration, reproducibility, and data pipelines."
        ),
    ],
    "todos": [
        (
            "Implement the full dual-loop agent framework.",
            "Complete BaseAMLAgent with fast-loop / slow-loop orchestration "
            "inheriting from StockSim's TraderAgent. Build the three role-"
            "specific agent classes. Implement role-specific strategy state "
            "dataclasses, observation context builder, local memory backend, "
            "strategy validator, and the LLM-shaped slow strategist pipeline."
        ),
        (
            "Build the strategy layer and LLM integration point.",
            "Implement the LLM strategist with structured context building, "
            "JSON response parsing, field whitelisting to prevent hallucination, "
            "and dataclass-safe update application. Use static test doubles "
            "to prove the control flow before wiring real LLM APIs."
        ),
        (
            "Harden the launcher and orchestration layer.",
            "Replace hardcoded time.sleep() waits with health-check polling, "
            "add process join timeouts with automatic termination of hung "
            "processes, fix signal handling for graceful shutdown with correct "
            "exit codes, and add a RabbitMQ TCP preflight check."
        ),
        (
            "Build context and memory infrastructure.",
            "Implement ObservationProcessor for structured market/portfolio/"
            "strategy context, LocalAgentMemory for in-process event storage, "
            "and AgentProfile dataclasses for stable behavioural identity. "
            "Define the MemoryBackend protocol for future Zep semantic memory."
        ),
    ],
    "challenges": [
        (
            "Understanding StockSim's architecture before extending it.",
            "StockSim is a substantial codebase with its own abstractions for "
            "exchange agents (order-book and candle-based), TraderAgent base "
            "class, RabbitMQ messaging protocols, simulation clock, data clients "
            "(Polygon, Alpha Vantage), and chart/report utilities. Reading and "
            "understanding the full architecture before building the AML-Sim "
            "layer on top of it required significant time investment. The "
            "submodule approach isolates changes but adds complexity to import "
            "paths and debugging."
        ),
        (
            "RabbitMQ unavailability on Windows development machine.",
            "The development machine (Windows 11 Home, Git Bash) does not have "
            "Docker Desktop or a native RabbitMQ installation. Without RabbitMQ, "
            "all inter-component communication fails — exchange agents, trader "
            "agents, and the simulation clock cannot coordinate. The dry-run mode "
            "(aml_runner.py --dry-run) validates everything up to the process "
            "launch point, but full integration testing is blocked."
        ),
        (
            "Designing the LLM strategist interface for safety and generality.",
            "Drawing from the literature — particularly the overfitting issues "
            "in Tan Shannon's dual-agent semantic loop (Mean IC 0.045 vs. "
            "syntax-only variant Mean IC 0.0454) — the slow loop must accept "
            "strategy proposals from arbitrary sources (LLM, rules engine, RL "
            "policy) while preventing dangerous parameter values. Designing the "
            "validator to catch not just syntax errors but semantically unsafe "
            "proposals required careful consideration of bounds and allowed fields."
        ),
    ],
    "plans": [
        (
            "Deep-dive into StockSim source code.",
            "Read all StockSim modules systematically — exchange agent, "
            "TraderAgent base class, simulation clock, order types (Market, "
            "Limit), message protocols, portfolio accounting — to build a "
            "complete mental model before extending it with AML-specific agents."
        ),
        (
            "Set up RabbitMQ via Docker Desktop or CloudAMQP.",
            "Install Docker Desktop for Windows (requires WSL2 enablement), "
            "then use StockSim's docker-compose.yml to start RabbitMQ with "
            "`docker compose up -d rabbitmq`. Alternatively, provision a free "
            "CloudAMQP instance for development. This unblocks all integration "
            "testing of the multi-process orchestration."
        ),
        (
            "Use dry-run mode for continuous validation.",
            "While RabbitMQ is unavailable, use `aml_runner.py --dry-run` to "
            "validate scenario parsing, config generation, run directory creation, "
            "and metadata writing. This covers the entire orchestration layer "
            "and ensures the code is always in a runnable state."
        ),
    ],
}

# =============================================================================
# JUNE Content
# =============================================================================
june_content = {
    "achievements": [
        (
            "Completed the full dual-loop agent framework.",
            "Implemented BaseAMLAgent (base.py) with shared fast-loop / "
            "slow-loop orchestration on top of StockSim's TraderAgent. Built "
            "three role-specific agent classes: AMLMarketMakerTrader "
            "(continuous bid/ask quoting with inventory skew), "
            "AMLRetailTrader (probabilistic noisy market orders with buy/sell "
            "bias), and AMLInstitutionalTrader (sliced child orders toward "
            "target positions). All run as independent multiprocessing.Process "
            "instances communicating through RabbitMQ."
        ),
        (
            "Built the strategy layer and LLM integration point.",
            "Defined role-specific strategy state dataclasses (MarketMaker, "
            "Retail, Institutional strategy states) inheriting from "
            "BaseStrategyState. Implemented the complete LLM strategist "
            "pipeline (LLMStrategist) with structured context building, JSON "
            "response parsing, strategy field whitelisting to prevent "
            "hallucination, and dataclass-safe update application. A "
            "StaticJSONLLMClient test double returns fixed role-specific "
            "responses, proving the control flow before wiring real LLM APIs."
        ),
        (
            "Implemented the strategy validation pipeline.",
            "Built a bounds validator (validator.py) that checks every "
            "strategy proposal before application: trade_probability and "
            "buy_bias must be in [0,1]; quote_size, spread, and "
            "child_order_size must be positive and within limits; confidence "
            "must be >=0; risk_mode must be one of {conservative, normal, "
            "aggressive}. Invalid proposals are rejected and logged; the "
            "agent keeps its previous state, ensuring a secure LLM pipeline."
        ),
        (
            "Built the context and memory infrastructure.",
            "Implemented ObservationProcessor that builds a structured context "
            "package at each tick containing agent identity, profile, current "
            "time, market snapshot, portfolio state (cash, value, per-"
            "instrument inventory, PnL), pending orders, recent fills, "
            "current strategy state, and memory. Built LocalAgentMemory for "
            "in-process event storage and a ZepAgentMemory placeholder for "
            "future semantic memory. Defined AgentProfile dataclasses with "
            "role-specific behavioural traits and a YAML-to-profile factory."
        ),
        (
            "Hardened the launcher and orchestration layer.",
            "Fixed the SIGINT handler (was raising SystemExit(0), now uses a "
            "flag and exits with code 130). Replaced hardcoded time.sleep() "
            "waits with health-check polling. Added per-process join timeout "
            "(300s) with automatic termination of hung processes. Added "
            "RabbitMQ TCP preflight check with clear error messages. Migrated "
            "all print() calls to structured logging. Made random seeds "
            "deterministic for reproducibility. Made the output-directory "
            "monkey-patch idempotent."
        ),
        (
            "Conducted full-system code review and applied fixes.",
            "Reviewed all ~2,100 lines across 40+ source files, identifying "
            "and fixing 24 issues: 5 critical (SIGINT, race conditions, "
            "timeout-free joins), 5 high (agent crash resilience, "
            "reproducibility, validation safety), 8 medium (serialization "
            "duplication, logging consistency, missing tests), and 6 low "
            "(dead code, import style, cosmetic). Created shared "
            "serialization module to eliminate duplicated logic across 3 "
            "files. Fixed non-dataclass mutation bug in the LLM strategist. "
            "Added hasattr guards to the validator."
        ),
        (
            "Prepared first interim report and monthly logs.",
            "Wrote the NUS First Interim Report covering project title, full "
            "scope and methodology, and a detailed 3-month development plan "
            "(July–September 2026) with concrete deliverables for each month. "
            "Established the four-phase development roadmap: Phase 1 — platform "
            "stabilisation and real LLM API integration (July), Phase 2 — "
            "Controlled Uncertainty Lab with ShockInjector infrastructure "
            "(August), Phase 3 — Human-in-the-Loop mode and live dashboard "
            "(September), Phase 4 — full experiments and final report "
            "(October–November). Authored detailed monthly logs for May and "
            "June documenting the full development trajectory from literature "
            "review through to production hardening."
        ),
        (
            "Connected literature insights to architectural decisions.",
            "Mapped concepts from the May literature review onto concrete "
            "implementation choices: the FinVision reflection module inspired "
            "our slow-loop strategist design (periodic strategy review with "
            "structured observation context); Tan Shannon's Syntax-Fix and "
            "Critique-Improvement agent loop informed our validator design "
            "(all proposals checked before application) and the field-"
            "whitelisting mechanism in the LLM strategist; the generation-"
            "evaluation split in Chen & Kawashima's two-agent framework "
            "validated our separation of strategist (proposer) from validator "
            "(evaluator). The literature's finding that semantic refinement "
            "can induce in-sample overfitting (Tan Shannon, Mean IC 0.045 vs. "
            "syntax-only 0.0454) reinforced the importance of conservative "
            "bounds in our validator."
        ),
    ],
    "todos": [
        (
            "Wire real LLM API clients (OpenAI / Anthropic).",
            "Replace StaticJSONLLMClient with real API adapters implementing "
            "the JSONLLMClient protocol. Test that all three agent roles "
            "receive and apply coherent strategy updates from actual LLM "
            "responses. Implement prompt engineering for financial decision-"
            "making context."
        ),
        (
            "Agent benchmarking and post-run analytics.",
            "Implement proper benchmarking for all agent roles across multi-"
            "hour simulation runs. Build post-run analysis scripts comparing "
            "performance metrics (Sharpe ratio, PnL, recovery time, inventory "
            "stability) across multiple runs."
        ),
        (
            "Wire the Zep memory backend.",
            "Replace the NotImplementedError placeholder in ZepAgentMemory "
            "with a working Zep SDK integration, enabling semantic memory "
            "retrieval across slow-loop invocations."
        ),
        (
            "Begin Controlled Uncertainty Lab design.",
            "Design the ShockInjector component, shock event taxonomy (price "
            "shocks, volatility spikes, information delays, fake news, "
            "liquidity crashes, spoofing), and the RabbitMQ topic topology "
            "for broadcasting shock events to agents."
        ),
    ],
    "challenges": [
        (
            "RabbitMQ unavailability blocking integration tests.",
            "Docker Desktop is not installed on the development Windows 11 "
            "machine, and there is no native RabbitMQ installation. All "
            "integration testing (full scenario runs with live processes) is "
            "blocked. The dry-run mode validates everything up to the process "
            "launch point, but the multi-process coordination cannot be "
            "verified end-to-end."
        ),
        (
            "Multi-process debugging complexity.",
            "With up to 9 independent processes (1 exchange + 7 traders + 1 "
            "clock), debugging is inherently difficult. A crash in one "
            "process (e.g., a trader agent encountering an unhandled "
            "exception in its slow loop) can silently cause that agent to "
            "disappear while the rest of the simulation continues, producing "
            "misleading results. The launcher now has health checks and "
            "timeouts, but agent-internal error resilience was a key focus "
            "of this month's hardening work."
        ),
        (
            "Balancing generality with safety in the LLM pipeline.",
            "The slow loop must accept strategy proposals from arbitrary "
            "sources (LLM, rules, RL) while preventing dangerous parameter "
            "values. The current validator checks numeric bounds and enum "
            "values, but more sophisticated checks (e.g., consistency between "
            "related fields) may be needed once real LLM outputs are tested."
        ),
    ],
    "plans": [
        (
            "Install Docker Desktop for Windows.",
            "Enable WSL2 and install Docker Desktop to run RabbitMQ via "
            "`docker compose up -d rabbitmq` from the StockSim directory. "
            "This unblocks all integration testing and is the recommended "
            "approach in the project README."
        ),
        (
            "Implement comprehensive agent error handling.",
            "Wrap every agent entry point (slow loop, fast loop, message "
            "handlers) in try/except blocks that log errors and allow the "
            "agent to continue operating with its last-known-good state. "
            "Add process-level health monitoring so the launcher detects "
            "and reports crashed agents."
        ),
        (
            "Add integration tests and scenario validation tests.",
            "Write pytest-based tests for: scenario YAML loading and "
            "validation, strategy validator boundary conditions, observation "
            "context builder output shape, serialization round-tripping, and "
            "LLM strategist response parsing. These can run without RabbitMQ "
            "and provide fast feedback during development."
        ),
        (
            "Test LLM outputs against the validator.",
            "Once real LLM API clients are wired, run batch tests to "
            "characterise how often LLM proposals fail validation. Use this "
            "data to tune prompts, adjust validator bounds, or add "
            "consistency checks. The goal is a pipeline where the LLM "
            "proposes, the validator catches unsafe proposals, and the "
            "agent degrades safely."
        ),
    ],
}

# =============================================================================
# Generate both files
# =============================================================================

may_doc = generate_monthly_log("May", "2026", "1 May 2026 – 31 May 2026", may_content)
may_path = OUT_DIR / "Monthly_Log_May_2026_PAN_YUFAN.docx"
may_doc.save(str(may_path))
print(f"Saved: {may_path}")

jun_doc = generate_monthly_log("June", "2026", "1 June 2026 – 26 June 2026", june_content)
jun_path = OUT_DIR / "Monthly_Log_June_2026_PAN_YUFAN.docx"
jun_doc.save(str(jun_path))
print(f"Saved: {jun_path}")

print("Done.")
