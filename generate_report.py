"""Generate NUS First Interim Report document for AML-Sim capstone project."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# --- Page style ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# --- Header ---
h = doc.sections[0].header
hp = h.paragraphs[0]
hp.text = "NUS School of Computing — First Interim Report"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.style.font.size = Pt(8)
hp.style.font.color.rgb = RGBColor(128, 128, 128)

# =============================================================================
# Title
# =============================================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "First Interim Report\n"
    "FT5007 FinTech Capstone Project"
)
run.font.size = Pt(16)
run.bold = True

doc.add_paragraph()

# Student info table
table = doc.add_table(rows=4, cols=2, style="Light Shading Accent 1")
cells = [
    ("Student Name", "PAN YUFAN"),
    ("Student ID", "[Your Student ID]"),
    ("Role", "Systems + Infrastructure Lead"),
    ("Academic Advisor", "[Advisor Name]"),
]
for i, (label, value) in enumerate(cells):
    table.cell(i, 0).text = label
    table.cell(i, 1).text = value
    for cell in table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)

doc.add_paragraph()

# =============================================================================
# Question 1
# =============================================================================
q1 = doc.add_heading("1. Project Title", level=1)

title_para = doc.add_paragraph()
run = title_para.add_run(
    "AML-Sim: A Multi-Agent Market Simulation Platform for Studying "
    "Adaptive Decision-Making Under Uncertainty"
)
run.bold = True
run.font.size = Pt(12)

# =============================================================================
# Question 2
# =============================================================================
doc.add_heading("2. Project Scope and Methodology", level=1)

doc.add_heading("2.1 Scope: The Four Systems", level=2)

doc.add_paragraph(
    "The project builds AML-Sim — a research-grade simulation platform that "
    "creates synthetic financial markets to study how decision-makers (both AI "
    "agents and humans) behave under controlled uncertainty. The platform "
    "consists of four integrated systems:"
)

systems = [
    (
        "Market Ecosystem Simulator",
        "A realistic exchange with a limit order book, liquidity provision "
        "through market makers, price discovery, and multiple interacting "
        "market participants — essentially a miniaturised Wall Street ecosystem "
        "running inside a single machine."
    ),
    (
        "AI Multi-Agent Environment",
        "Heterogeneous trading agents (market makers, retail traders, "
        "institutional traders, and — in future iterations — manipulators and "
        "regulators), each with different information sets, goals, latency "
        "profiles, and risk appetites. Their continuous interaction creates "
        "emergent market behaviour."
    ),
    (
        "Controlled Uncertainty Lab",
        "Infrastructure to inject shocks at precise simulation times: fake "
        "news events, volatility spikes, information delays, liquidity "
        "crashes, adversarial spoofing, and policy interventions. The "
        "platform then observes who adapts, who collapses, and who stabilises "
        "fastest."
    ),
    (
        "Human Decision Experiment Platform",
        "A human-in-the-loop mode where real participants trade inside the "
        "ecosystem alongside AI agents. The system studies panic reactions, "
        "overconfidence, recovery ability, risk calibration, and adaptation "
        "speed — connecting to behavioural finance, cognitive science, and "
        "adaptive expertise research."
    ),
]

for name, desc in systems:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}. ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "The deep research question driving the project is: What makes a "
    "decision-maker robust under uncertainty? The platform measures "
    "resilience, adaptability, stability, calibration, and recoverability — "
    "not raw profit."
)

doc.add_heading("2.2 Role: Systems + Infrastructure Lead", level=2)

doc.add_paragraph(
    "My responsibility is to build the 'world' that agents and experiments "
    "run inside. Specifically:"
)

responsibilities = [
    "Exchange engine and order-book integration (via StockSim submodule).",
    "System architecture: fast-loop / slow-loop agent design, component "
    "boundaries, separation of concerns between StockSim and AML-Sim.",
    "Concurrency model: multiprocess orchestration for exchange agents, "
    "trader agents, and the simulation clock, coordinated through "
    "RabbitMQ pub/sub messaging.",
    "Networking: RabbitMQ message bus topology, agent-to-exchange "
    "communication protocols, message serialisation.",
    "Simulation orchestration: process lifecycle (start, health-check, "
    "graceful shutdown), signal handling, run directory management.",
    "Reproducibility: deterministic seeding, configuration archiving, "
    "run-specific artifact directories, metadata tracking.",
    "Data pipelines: market data ingestion, synthetic data generation, "
    "trader action ledgers, post-simulation report generation.",
]

for r in responsibilities:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("2.3 Methodology", level=2)

doc.add_paragraph(
    "The platform builds on StockSim, an open-source market simulator that "
    "provides the core exchange engine, order matching, and a RabbitMQ-based "
    "message bus. AML-Sim wraps StockSim as a git submodule and extends it "
    "with several architectural layers:"
)

method_items = [
    (
        "Fast-loop / slow-loop agent architecture.",
        "Each agent runs a high-frequency execution loop (submitting orders, "
        "cancelling stale quotes) driven by a current strategy state, and a "
        "lower-frequency strategic loop that proposes updates to the strategy "
        "state based on market observations. The slow loop is designed as an "
        "LLM-integration point — strategy proposals can come from an LLM, a "
        "rule-based engine, or reinforcement learning policies. All proposals "
        "pass through a bounds validator before being applied."
    ),
    (
        "Multiprocess orchestration.",
        "Exchange agents, trader agents, and the simulation clock each run in "
        "independent Python processes, coordinated through RabbitMQ pub/sub "
        "messaging. The AML launcher manages process lifecycle with health "
        "checks, timeout-based joins, and graceful termination on shutdown "
        "signals."
    ),
    (
        "Deterministic reproducibility.",
        "Every run archives its full scenario configuration, generates a "
        "run-specific artifact directory under .aml_runs/, and uses "
        "deterministic seeding derived from run identity. This enables "
        "controlled A/B experiments — identical market and agents, different "
        "shock configurations."
    ),
    (
        "Strategy validation pipeline.",
        "All strategy state updates pass through a bounds validator before "
        "application. The validator checks trade probability, buy bias, "
        "quote size, spread, child order size, confidence, and risk mode "
        "against configurable limits. Invalid proposals are rejected and "
        "logged; the agent keeps its previous state."
    ),
    (
        "Observation context builder.",
        "At each simulation tick, a structured snapshot is built containing "
        "market state, portfolio state, pending orders, recent fills, current "
        "strategy state, and memory events. This is the input package for "
        "the slow-loop strategist — and, in future work, for the LLM prompt."
    ),
    (
        "Iterative development approach.",
        "Phase 1 proves the engine and agent loops with synthetic data and "
        "static strategy policies. Phase 2 wires in real LLM API calls. "
        "Phase 3 adds the uncertainty injection infrastructure. Phase 4 "
        "adds human participation. Each phase is tested end-to-end before "
        "the next begins."
    ),
]

for label, desc in method_items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(desc)

# =============================================================================
# Question 3
# =============================================================================
doc.add_heading("3. Plan for the Next Three Months (July – September 2026)", level=1)

doc.add_heading("Month 1 (July 2026): Stabilise the Platform and Validate Agent Readiness", level=2)

month1 = [
    (
        "Wire real LLM API clients.",
        "Replace the current StaticJSONLLMClient (which returns hard-coded "
        "JSON responses for testing) with real LLM API adapters (OpenAI / "
        "Anthropic) implementing the JSONLLMClient protocol. This proves the "
        "LLM-shaped control flow with actual model responses."
    ),
    (
        "Agent benchmarking and validation.",
        "Implement proper benchmarking for all three agent roles (market "
        "maker, retail, institutional) — verifying that the fast/slow loops, "
        "portfolio accounting, and action ledgers work correctly end-to-end "
        "across multi-hour simulation runs."
    ),
    (
        "Post-run analytics.",
        "Add basic post-run analysis scripts that compare agent performance "
        "metrics (Sharpe ratio, PnL, recovery time, inventory stability) "
        "across multiple runs."
    ),
    (
        "Memory backend integration.",
        "Wire the ZepAgentMemory backend (currently a NotImplementedError "
        "placeholder) so that agents can retain semantic memory across slow-loop "
        "invocations instead of the current simple local event store."
    ),
]

for label, desc in month1:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Month 2 (August 2026): Build the Controlled Uncertainty Lab (System 3)", level=2)

month2 = [
    (
        "ShockInjector infrastructure.",
        "Design and implement a ShockInjector component that runs as a "
        "separate process and broadcasts shock events over RabbitMQ at "
        "precise simulation-clock-coordinated times."
    ),
    (
        "Shock taxonomy.",
        "Define and implement shock patterns: price shocks, volatility "
        "spikes, information delays, fake news injections, liquidity "
        "crashes, and adversarial spoofing. Each shock type has configurable "
        "parameters (magnitude, duration, affected instruments)."
    ),
    (
        "Agent shock handling.",
        "Modify agent message handlers to receive and react to shock events. "
        "For example: market makers widen spreads after volatility shocks; "
        "retail traders increase herding tendency after fake news; "
        "institutional traders pause execution during circuit-breaker events."
    ),
    (
        "Experiment runner.",
        "Build an experiment runner that performs parameter sweeps across "
        "multiple shock configurations and generates comparative reports "
        "across runs. This is the infrastructure that enables the core "
        "research question."
    ),
]

for label, desc in month2:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Month 3 (September 2026): Human-in-the-Loop and Initial Experiments", level=2)

month3 = [
    (
        "Human-in-the-loop (HITL) mode.",
        "Build the infrastructure for System 4: simulation pause/resume "
        "mechanisms, human order submission via REST or WebSocket API, "
        "and recording of human actions in the same ledger schema as "
        "agent actions."
    ),
    (
        "Live dashboard.",
        "Build a dashboard foundation for real-time monitoring: order-book "
        "visualisation, agent PnL streaming, shock event log, and a simple "
        "web-based interface for human participants."
    ),
    (
        "First controlled experiments.",
        "Run the first batch of controlled experiments with AI agents under "
        "no-shock vs. shock conditions. Document preliminary findings about "
        "what strategies produce robust decision-making under uncertainty."
    ),
    (
        "Infrastructure for final report.",
        "Prepare the infrastructure for the final report: ensure end-to-end "
        "reproducible runs, clear architecture diagrams, documented experiment "
        "configurations, and publication-quality output formats."
    ),
]

for label, desc in month3:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

# --- Footer ---
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("— End of First Interim Report —")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.italic = True

# =============================================================================
# Save
# =============================================================================
output_path = (
    r"C:\Users\mrp\Desktop\Capstone_public\AML-Sim\docs"
    r"\First_Interim_Report_PAN_YUFAN.docx"
)
import os
from pathlib import Path
Path(output_path).parent.mkdir(parents=True, exist_ok=True)
doc.save(output_path)
print(f"Saved: {output_path}")
