"""Generate May and June 2026 monthly log .docx files for NUS capstone."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path.home() / "Desktop"


def add_header(doc: Document, month: str) -> None:
    """Common header for a monthly log."""
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Capstone Project Monthly Log — {month} 2026")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=3, cols=2, style="Light Shading Accent 1")
    cells = [
        ("Student Name", "PAN YUFAN"),
        ("Project", "AML-Sim: A Multi-Agent Market Simulation Platform"),
        ("Role", "Systems + Infrastructure Lead"),
    ]
    for i, (label, value) in enumerate(cells):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)

    doc.add_paragraph()


def add_section(doc: Document, title: str, body: str) -> None:
    """Add a titled section with body text."""
    h = doc.add_heading(title, level=2)
    doc.add_paragraph(body)


def add_bullet_section(doc: Document, title: str, items: list[tuple[str, str]]) -> None:
    """Add a section with bold-label + description bullets."""
    doc.add_heading(title, level=2)
    for label, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(desc)


def add_signoff(doc: Document) -> None:
    """Standard footer block."""
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Monthly Log —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True


# =============================================================================
# MAY 2026
# =============================================================================

doc_may = Document()
style = doc_may.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

add_header(doc_may, "May")

add_section(
    doc_may,
    "Overview",
    "May marked the project kick-off. The primary focus was on literature "
    "review, understanding the research landscape, defining the system "
    "architecture, and performing initial exploratory coding to validate "
    "the technology stack.",
)

add_bullet_section(
    doc_may,
    "Literature Review",
    [
        (
            "Market simulation and agent-based modelling.",
            "Surveyed existing market simulation platforms (StockSim, ABIDES, "
            "JADE, MASON) to understand the state of the art in multi-agent "
            "financial simulation. Identified StockSim as the most suitable "
            "base engine due to its RabbitMQ-based component architecture, "
            "clean TraderAgent abstraction, and YAML-driven configuration."
        ),
        (
            "Behavioural finance and adaptive expertise.",
            "Reviewed literature on decision-making under uncertainty, "
            "including Kahneman & Tversky's prospect theory, Gigerenzer's "
            "ecological rationality, and research on adaptive expertise by "
            "Hatano & Inagaki. This shaped the project's core research "
            "question: what makes a decision-maker robust under uncertainty?"
        ),
        (
            "LLM-based agents in economics.",
            "Studied recent work on LLM-powered trading agents (Horton's "
            "Large Language Models as Simulated Economic Agents, Park et al.'s "
            "Generative Agents) to understand prompt engineering strategies, "
            "multi-step reasoning for financial decisions, and hallucination "
            "mitigation in quantitative contexts."
        ),
        (
            "Uncertainty injection in simulations.",
            "Reviewed methods for introducing controlled shocks in agent-based "
            "models — price shocks, information asymmetry, latency variation, "
            "and adversarial perturbation — to inform the design of System 3 "
            "(Controlled Uncertainty Lab)."
        ),
    ],
)

add_bullet_section(
    doc_may,
    "Architecture Discussion & Design",
    [
        (
            "Four-system decomposition.",
            "Worked with the project advisor and teammate to decompose the "
            "AML platform into four integrated systems: (1) Market Ecosystem "
            "Simulator, (2) AI Multi-Agent Environment, (3) Controlled "
            "Uncertainty Lab, and (4) Human Decision Experiment Platform. "
            "Agreed that Systems 1–2 would be built first as a foundation."
        ),
        (
            "Two-layer architecture decision.",
            "Decided to keep AML-Sim and StockSim as separate layers: "
            "StockSim owns execution, messaging, portfolio accounting, and "
            "order state; AML-Sim owns agent behaviour, strategy, memory, "
            "observation packaging, and experiment orchestration. This "
            "separation of concerns allows StockSim to evolve independently "
            "as an upstream dependency."
        ),
        (
            "Fast-loop / slow-loop agent design.",
            "Designed the dual-loop agent architecture: a high-frequency "
            "execution loop for order placement (role-specific) and a "
            "lower-frequency strategic loop for policy updates (LLM-driven). "
            "The slow loop always passes through a strategy validator before "
            "any proposal is applied, ensuring safety bounds even with LLM "
            "output."
        ),
        (
            "Technology stack confirmation.",
            "Confirmed the technology stack: Python 3.13 for all components, "
            "RabbitMQ for inter-process messaging, multiprocessing for "
            "component isolation, asyncio for agent-internal concurrency, "
            "YAML for scenario definition, and JSON for action ledgers and "
            "reports. Git submodules for StockSim dependency management."
        ),
    ],
)

add_bullet_section(
    doc_may,
    "Initial Code Exploration",
    [
        (
            "StockSim submodule integration.",
            "Cloned the StockSim repository as a git submodule, installed "
            "its Python dependencies, and verified that the standalone "
            "StockSim demo scenarios run correctly."
        ),
        (
            "Proof-of-concept AML runner.",
            "Wrote an initial version of aml_runner.py that reads an AML "
            "scenario YAML file, extracts and validates the StockSim config "
            "section, creates timestamped run directories, and archives "
            "artifacts for reproducibility."
        ),
        (
            "Synthetic order-book scenario.",
            "Created the first AML scenario (aml_orderbook_replay.yaml) "
            "using a 30-minute synthetic AAPL order book with one market "
            "maker, five retail traders, and one institutional trader. "
            "Verified that the scenario parses correctly and generates valid "
            "StockSim configuration."
        ),
        (
            "Agent class skeleton.",
            "Began implementing the three AML agent classes as subclasses of "
            "StockSim's TraderAgent, establishing the inheritance chain and "
            "testing that the component launcher can start them as independent "
            "processes."
        ),
    ],
)

add_signoff(doc_may)

may_path = OUT_DIR / "Monthly_Log_May_2026_PAN_YUFAN.docx"
doc_may.save(str(may_path))
print(f"Saved: {may_path}")

# =============================================================================
# JUNE 2026
# =============================================================================

doc_jun = Document()
style2 = doc_jun.styles["Normal"]
font2 = style2.font
font2.name = "Calibri"
font2.size = Pt(11)
style2.paragraph_format.space_after = Pt(6)
style2.paragraph_format.line_spacing = 1.15

add_header(doc_jun, "June")

add_section(
    doc_jun,
    "Overview",
    "June focused on building the core platform infrastructure. The agent "
    "framework was completed with the fast-loop / slow-loop architecture, "
    "the launcher was hardened for production reliability, the codebase "
    "underwent rigorous review, and the first interim report was prepared.",
)

add_bullet_section(
    doc_jun,
    "Agent Framework Implementation",
    [
        (
            "BaseAMLAgent with dual-loop architecture.",
            "Implemented the shared base class (base.py) that all three AML "
            "agent roles inherit from. It provides the fast-loop / slow-loop "
            "orchestration: each simulation tick, the agent checks whether "
            "the slow loop is due, runs the strategist if so, validates the "
            "proposal, rebuilds the observation, then runs the role-specific "
            "fast loop. The base class also owns action event recording, "
            "portfolio snapshotting, and serialisation."
        ),
        (
            "Market maker agent (market_maker_trader.py).",
            "Implements a continuous quoting strategy: on each fast-loop "
            "tick, it cancels outstanding quotes and posts new bid/ask limit "
            "orders around a fair price with configurable spread, quote size, "
            "target inventory, and inventory skew. The skew adjusts the "
            "midpoint to prevent unbounded inventory accumulation."
        ),
        (
            "Retail trader agent (retail_trader.py).",
            "Models a small noisy participant: each fast-loop tick it "
            "probabilistically submits a market order of random size (up to "
            "a configured maximum) with a configurable buy/sell bias. "
            "Inventory-aware: will not sell shares it does not hold."
        ),
        (
            "Institutional trader agent (institutional_trader.py).",
            "Models a larger participant working toward target positions "
            "through sliced child orders. Each fast-loop tick it places one "
            "child order toward the target, using either MARKET or LIMIT "
            "order types. The agent tracks the gap between current holdings "
            "and the target and stops when the position is reached."
        ),
    ],
)

add_bullet_section(
    doc_jun,
    "Strategy Layer & Validation",
    [
        (
            "Strategy state models (state.py).",
            "Defined role-specific dataclasses for strategy state: "
            "MarketMakerStrategyState (fair_price, spread, quote_size, "
            "target_inventory, inventory_skew), RetailStrategyState "
            "(trade_probability, buy_bias, max_order_size, herding_tendency, "
            "panic_level), and InstitutionalStrategyState (target_positions, "
            "child_order_size, order_type, execution_style, urgency). All "
            "inherit from BaseStrategyState (risk_mode, confidence, reason, "
            "updated_at)."
        ),
        (
            "LLM-shaped slow strategist (llm_slow_strategy.py).",
            "Built the complete LLM strategist pipeline: build structured "
            "context from observation + profile + memory, call an LLM client "
            "via the JSONLLMClient protocol, parse the JSON response to "
            "extract strategy_updates + confidence + reason, and apply "
            "updates through dataclasses.replace(). A StaticJSONLLMClient "
            "test double returns fixed role-specific responses, proving the "
            "control flow before wiring real LLM APIs."
        ),
        (
            "Strategy validator (validator.py).",
            "Implemented bounds checking that runs on every strategy proposal "
            "before it is applied. Checks: trade_probability and buy_bias "
            "must be in [0, 1]; quote_size must be >0 and <=10,000; spread "
            "must be >0; child_order_size must be >0 and <=100,000; "
            "confidence must be >=0; risk_mode must be one of {conservative, "
            "normal, aggressive}. Invalid proposals are rejected and logged; "
            "the agent keeps its previous strategy state."
        ),
    ],
)

add_bullet_section(
    doc_jun,
    "Context & Memory Infrastructure",
    [
        (
            "Observation processor (observation.py).",
            "Implemented a structured context builder that produces the "
            "input package for the slow-loop strategist at each tick. The "
            "package includes agent identity + profile, current simulation "
            "time, latest market snapshot, cash, portfolio value, per-"
            "instrument inventory (long, short, net, last_price, realized_pnl), "
            "pending orders, recent fills, current strategy state, and "
            "memory context."
        ),
        (
            "Local memory backend (memory.py).",
            "Implemented an in-process memory store that records agent events "
            "and retrieves the most recent N events as context for the slow "
            "loop. Defined a MemoryBackend protocol shared with a future Zep "
            "semantic-memory backend (placeholder implemented, ready to wire "
            "when the Zep SDK is available)."
        ),
        (
            "Agent profile models (profile.py).",
            "Defined stable behavioural identity dataclasses: base "
            "AgentProfile with role, risk_tolerance, decision_style, "
            "personality, behavioural_traits, and preferences; plus role-"
            "specific subclasses (MarketMakerProfile, RetailProfile, "
            "InstitutionalProfile) with role-appropriate defaults. A "
            "coerce_profile() factory converts YAML dicts to typed profiles, "
            "isolating unrecognised fields under a safe namespace."
        ),
    ],
)

add_bullet_section(
    doc_jun,
    "Launcher & Orchestration Hardening",
    [
        (
            "Signal handling fix.",
            "Replaced a SystemExit(0) call inside the SIGINT handler with a "
            "flag-based shutdown mechanism, ensuring the orchestrator exits "
            "with the correct code (130) and still generates post-simulation "
            "reports after an interrupt."
        ),
        (
            "Process health checking.",
            "Replaced hardcoded time.sleep(10)/time.sleep(20) waits with "
            "polling loops that verify all exchange and trader processes are "
            "alive before starting the simulation clock, with configurable "
            "deadlines."
        ),
        (
            "Process join timeout.",
            "Added per-process join timeout (300 s total) with automatic "
            "termination of hung processes instead of blocking forever."
        ),
        (
            "RabbitMQ preflight check.",
            "Added a TCP health check at launch time that fails early with "
            "a clear error message if RabbitMQ is not reachable, instead of "
            "crashing deep inside aio-pika imports."
        ),
        (
            "Logging migration.",
            "Replaced all print() calls in the launcher with structured "
            "logging (logging.getLogger), ensuring that orchestration "
            "messages appear in run logs alongside agent messages."
        ),
        (
            "Reproducibility.",
            "Made Random_Trader seeds deterministic by deriving them from "
            "(run_id, agent_id) instead of time-based random.randint(). "
            "Made the StockSim output-directory monkey-patch idempotent "
            "with a sentinel guard."
        ),
    ],
)

add_bullet_section(
    doc_jun,
    "Code Quality & Review",
    [
        (
            "Full-system code review.",
            "Performed a rigorous review of all ~2,100 lines of code across "
            "40+ source files, identifying 24 issues across critical, high, "
            "medium, and low severity tiers. Every finding was addressed."
        ),
        (
            "Shared serialization module (serialization.py).",
            "Extracted duplicated _serialize_value / _serialize_mapping "
            "functions from three files (base.py, observation.py, memory.py) "
            "into a single shared module. The unified serializer safely "
            "handles datetime, dataclass, Mapping, list, tuple, and now "
            "non-JSON-native types (set, bytes, Decimal, UUID) with a "
            "string fallback instead of crashing at json.dump time."
        ),
        (
            "Non-dataclass mutation fix.",
            "Fixed a correctness bug in the LLM strategist where the non-"
            "dataclass path of _apply_updates mutated the live strategy "
            "object in place via setattr before validation. Now uses "
            "copy.deepcopy to protect the original."
        ),
        (
            "Validator hardening.",
            "Added hasattr guards for confidence and risk_mode field access "
            "in the strategy validator so incomplete strategy state objects "
            "produce a clear error message instead of an AttributeError crash."
        ),
        (
            "Reporting improvements.",
            "Replaced the always-zero duration_days field with a human-"
            "readable _format_simulation_duration() that shows minutes, "
            "hours, or days depending on the simulation window. Moved "
            "the redundant parse_datetime_utc import to a proper helper."
        ),
    ],
)

add_bullet_section(
    doc_jun,
    "Documentation & Planning",
    [
        (
            "First interim report.",
            "Prepared the NUS First Interim Report covering the project "
            "title, full scope and methodology description, and a detailed "
            "3-month plan (July–September 2026) with concrete deliverables "
            "for each month."
        ),
        (
            "Project roadmap.",
            "Established the four-phase development roadmap: Phase 1 — "
            "platform stabilisation and real LLM API integration (July); "
            "Phase 2 — Controlled Uncertainty Lab with ShockInjector "
            "infrastructure (August); Phase 3 — Human-in-the-Loop mode "
            "and live dashboard (September); Phase 4 — full experiments "
            "and final report (October–November)."
        ),
    ],
)

add_signoff(doc_jun)

jun_path = OUT_DIR / "Monthly_Log_June_2026_PAN_YUFAN.docx"
doc_jun.save(str(jun_path))
print(f"Saved: {jun_path}")
